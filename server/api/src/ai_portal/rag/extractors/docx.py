"""DOCX extractor via :mod:`python-docx`."""
from __future__ import annotations

from io import BytesIO
from typing import Any

from ai_portal.rag.extractors.protocol import (
    Block,
    ExtractedDocument,
    HeadingBlock,
    ParagraphBlock,
    TableBlock,
)


class DocxExtractor:
    name = "docx"
    mime_types = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }

    def supports(self, mime: str) -> bool:
        return mime in self.mime_types

    async def extract(self, data: bytes, meta: dict[str, Any]) -> ExtractedDocument:
        from docx import Document  # type: ignore

        doc = Document(BytesIO(data))
        blocks: list[Block] = []
        text_parts: list[str] = []
        for para in doc.paragraphs:
            txt = (para.text or "").strip()
            if not txt:
                continue
            style = (getattr(para.style, "name", "") or "").lower()
            if style.startswith("heading"):
                try:
                    level = int(style.replace("heading", "").strip() or "1")
                except ValueError:
                    level = 1
                blocks.append(HeadingBlock(text=txt, level=level))
            else:
                blocks.append(ParagraphBlock(text=txt))
            text_parts.append(txt)
        for tbl in doc.tables:
            rows = [
                [cell.text.strip() for cell in row.cells]
                for row in tbl.rows
            ]
            if rows:
                blocks.append(TableBlock(rows=rows))
        title: str | None = None
        try:
            cp = doc.core_properties
            title = cp.title or None
        except Exception:
            pass
        return ExtractedDocument(
            text="\n".join(text_parts),
            blocks=blocks,
            meta={**meta, "title": title},
        )


__all__ = ["DocxExtractor"]
